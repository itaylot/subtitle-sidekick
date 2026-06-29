"""engine.py — transcription logic + save/export, fully decoupled from the UI.

faster-whisper + ivrit-ai model. Automatic GPU detection (NVIDIA → significantly faster).
"""

import os
import re
import json
import time
import shutil

MODEL_ACCURATE = "ivrit-ai/whisper-large-v3-turbo-ct2"  # Hebrew-specialized model
MODEL_FAST = "small"                                    # faster, less accurate
MAX_CUE_WORDS = 6  # max words per subtitle cue — improves on-screen readability

CPU_THREADS = max(4, os.cpu_count() or 4)

_model = None
_model_name = None
_device = "cpu"

# Cancel is handled by killing the worker process (see worker.py / app.py) — faster-whisper's
# C-level work can't be interrupted from within the same process. Pause is cooperative via a
# pause_check callback that blocks between segments.


def fmt_time(t: float) -> str:
    h = int(t // 3600); m = int((t % 3600) // 60); s = int(t % 60)
    ms = int(round((t - int(t)) * 1000))
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


# ── hardware detection ──
def _resolve_device():
    """Returns (device, compute_type). Uses NVIDIA GPU if available."""
    try:
        import ctranslate2
        if ctranslate2.get_cuda_device_count() > 0:
            return "cuda", "float16"
    except Exception:
        pass
    return "cpu", "int8"


def _load_model(name):
    """Load model; tries GPU first, falls back to CPU on failure."""
    from faster_whisper import WhisperModel
    dev, ct = _resolve_device()
    try:
        m = WhisperModel(name, device=dev, compute_type=ct, cpu_threads=CPU_THREADS, num_workers=1)
        return m, dev
    except Exception:
        m = WhisperModel(name, device="cpu", compute_type="int8", cpu_threads=CPU_THREADS, num_workers=1)
        return m, "cpu"


# ── SRT ──
def write_srt(srt_path, cues):
    with open(srt_path, "w", encoding="utf-8") as f:
        for i, c in enumerate(cues, 1):
            f.write(f"{i}\n{fmt_time(c['start'])} --> {fmt_time(c['end'])}\n{c['text'].strip()}\n\n")


def save_srt(video, cues):
    """Re-save SRT + viewer after editing."""
    srt = os.path.splitext(video)[0] + ".srt"
    write_srt(srt, cues)
    make_viewer(video, cues)
    return srt


# ── transcript export ──
def export_txt(video, cues):
    out = os.path.splitext(video)[0] + " — תמליל.txt"
    with open(out, "w", encoding="utf-8") as f:
        f.write("\n".join(c["text"].strip() for c in cues))
    return out


def export_docx(video, cues):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement

    def set_rtl(p):
        pPr = p._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc = Document()
    title = doc.add_heading(os.path.splitext(os.path.basename(video))[0], level=1)
    set_rtl(title)
    for c in cues:
        p = doc.add_paragraph(c["text"].strip())
        set_rtl(p)
    out = os.path.splitext(video)[0] + " — תמליל.docx"
    doc.save(out)
    return out


# standalone viewer. Native subtitles via <track> VTT — visible in fullscreen too.
VIEWER_TEMPLATE = """<!DOCTYPE html>
<html lang="he" dir="rtl"><head><meta charset="utf-8"><title>__TITLE__</title>
<style>
 body{margin:0;background:#15151a;font-family:system-ui,Arial,sans-serif}
 #stage{position:relative;max-width:1100px;margin:0 auto}
 video{width:100%;display:block;background:#000}
 video::cue{background:rgba(0,0,0,.72);color:#fff;font-size:1.05em;line-height:1.4}
</style></head><body>
<div id="stage"><video id="v" src="__SRC__" controls autoplay></video></div>
<script>
 const cues=__CUES__,v=document.getElementById('v');
 function ts(t){const h=String(Math.floor(t/3600)).padStart(2,'0'),
   m=String(Math.floor(t%3600/60)).padStart(2,'0'),s=String(Math.floor(t%60)).padStart(2,'0'),
   ms=String(Math.round(t%1*1000)).padStart(3,'0');return `${h}:${m}:${s}.${ms}`;}
 let vtt="WEBVTT\\n\\n";
 for(const c of cues){vtt+=ts(c.start)+" --> "+ts(c.end)+"\\n"+c.text+"\\n\\n";}
 const tr=document.createElement('track');
 tr.kind='subtitles';tr.srclang='he';tr.label='עברית';tr.default=true;
 tr.src=URL.createObjectURL(new Blob([vtt],{type:'text/vtt'}));
 v.appendChild(tr);
 v.addEventListener('loadedmetadata',()=>{if(v.textTracks[0])v.textTracks[0].mode='showing';});
</script></body></html>"""


def make_viewer(video, cues):
    folder = os.path.dirname(video)
    base = os.path.splitext(os.path.basename(video))[0]
    html = (VIEWER_TEMPLATE
            .replace("__TITLE__", base)
            .replace("__SRC__", os.path.basename(video))
            .replace("__CUES__", json.dumps(cues, ensure_ascii=False)))
    out = os.path.join(folder, base + " — כתוביות.html")
    with open(out, "w", encoding="utf-8") as f:
        f.write(html)
    return out


# ── lecture library: courses + transcribed lectures, stored as JSON ──
LIB_DIR = os.path.join(os.path.expanduser("~"), "Videos", "Subtitle Sidekick")
LIB_PATH = os.path.join(LIB_DIR, "library.json")


def load_library():
    """Returns {courses: [...], lectures: [...]} (creates empty if missing)."""
    try:
        with open(LIB_PATH, encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        data = {}
    data.setdefault("courses", [])
    data.setdefault("lectures", [])
    # filter out lectures whose file no longer exists on disk
    data["lectures"] = [l for l in data["lectures"] if os.path.exists(l.get("video", ""))]
    return data


def save_library(data):
    os.makedirs(LIB_DIR, exist_ok=True)
    with open(LIB_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return data


# ── user settings (transcription mode + personal cloud server config) — stored as local JSON ──
SETTINGS_PATH = os.path.join(LIB_DIR, "settings.json")


def load_settings():
    """Returns settings with all defaults filled in (transcription mode + cloud config and cost counter)."""
    try:
        with open(SETTINGS_PATH, encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        data = {}
    data.setdefault("transcription_mode", "local_accurate")
    data.setdefault("library_dir", LIB_DIR)  # base folder where course folders + lectures are stored
    data.setdefault("cloud", {})
    c = data["cloud"]
    c.setdefault("endpoint_url", "")
    c.setdefault("api_key", "")
    c.setdefault("price_per_hour", 0)        # GPU price per hour ($) — for cost calculation
    c.setdefault("total_seconds", 0)         # cumulative processing time on server
    c.setdefault("total_cost", 0)            # cumulative cost ($)
    return data


def save_settings(update):
    """Merges a partial update into existing settings (non-destructive) and saves. Returns full settings."""
    data = load_settings()
    for k, v in (update or {}).items():
        if k == "cloud" and isinstance(v, dict):
            data["cloud"].update(v)
        else:
            data[k] = v
    os.makedirs(LIB_DIR, exist_ok=True)
    with open(SETTINGS_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return data


def add_cloud_usage(seconds):
    """Accumulates server processing time into cumulative cost (based on configured price per hour)."""
    seconds = max(0, float(seconds or 0))
    data = load_settings()
    cloud = data["cloud"]
    cloud["total_seconds"] = float(cloud.get("total_seconds") or 0) + seconds
    cloud["total_cost"] = float(cloud.get("total_cost") or 0) + seconds / 3600.0 * float(cloud.get("price_per_hour") or 0)
    os.makedirs(LIB_DIR, exist_ok=True)
    with open(SETTINGS_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return data


# ── transcription queue: persisted jobs so the queue survives app restarts/crashes ──
QUEUE_PATH = os.path.join(LIB_DIR, "queue.json")


def library_dir():
    """User-chosen base folder for storing lectures (defaults to LIB_DIR)."""
    return load_settings().get("library_dir") or LIB_DIR


def _safe_folder(name):
    """Sanitize a course name into a valid Windows folder name."""
    return re.sub(r'[<>:"/\\|?*]', "_", (name or "").strip())


def _unique_path(path):
    """If path exists, return 'base (1).ext', 'base (2).ext'… so we never overwrite."""
    if not os.path.exists(path):
        return path
    base, ext = os.path.splitext(path)
    i = 1
    while os.path.exists(f"{base} ({i}){ext}"):
        i += 1
    return f"{base} ({i}){ext}"


def relocate(res, course):
    """Move a finished lecture's video into <library>/<course>/ and regenerate srt+viewer there.

    Returns res with updated paths. Empty course → unchanged (stays next to source).
    The viewer HTML embeds the video basename, so we regenerate srt+viewer from cues at the
    destination rather than moving stale copies. Raises on failure (caller marks the job failed).
    """
    folder = _safe_folder(course)
    if not folder:
        return res
    dest_dir = os.path.join(library_dir(), folder)
    os.makedirs(dest_dir, exist_ok=True)
    old_video, old_srt, old_viewer = res["video"], res.get("srt"), res.get("viewer")

    new_video = _unique_path(os.path.join(dest_dir, os.path.basename(old_video)))
    shutil.move(old_video, new_video)              # rename on same drive, copy across drives
    new_srt = os.path.splitext(new_video)[0] + ".srt"
    write_srt(new_srt, res["cues"])
    new_viewer = make_viewer(new_video, res["cues"])

    for p in (old_srt, old_viewer):                # drop now-stale in-place outputs
        if p and os.path.exists(p):
            try:
                os.remove(p)
            except OSError:
                pass

    res = dict(res)
    res["video"], res["srt"], res["viewer"] = new_video, new_srt, new_viewer
    return res


def load_queue():
    """Load the persisted queue for crash recovery. Returns only resumable jobs (status='queued').

    'running' jobs (interrupted mid-transcription) revert to 'queued'; 'done'/'failed' are history
    and dropped; jobs whose source file vanished are dropped (nothing to run).
    """
    try:
        with open(QUEUE_PATH, encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        return []
    jobs = data.get("jobs") if isinstance(data, dict) else None
    if not isinstance(jobs, list):
        return []
    out = []
    for j in jobs:
        if not isinstance(j, dict) or not j.get("sourcePath"):
            continue
        if j.get("status") not in ("queued", "running"):
            continue
        if not os.path.exists(j["sourcePath"]):
            continue
        j["status"] = "queued"
        j["error"] = None
        out.append(j)
    return out


def save_queue(jobs):
    """Atomically persist the queue (write to .tmp then os.replace) so a crash can't corrupt it."""
    os.makedirs(LIB_DIR, exist_ok=True)
    tmp = QUEUE_PATH + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump({"version": 1, "jobs": jobs or []}, f, ensure_ascii=False, indent=2)
    os.replace(tmp, QUEUE_PATH)
    return True


def create_course(name):
    name = (name or "").strip()
    data = load_library()
    if name and name not in data["courses"]:
        data["courses"].append(name)
        save_library(data)
    return data


def remove_course(name):
    """Remove a course; its lectures move to 'no course' (files are not deleted)."""
    data = load_library()
    data["courses"] = [c for c in data["courses"] if c != name]
    for l in data["lectures"]:
        if l.get("course") == name:
            l["course"] = ""
    return save_library(data)


def add_lecture(video, srt=None, course="", title=None):
    """Register a lecture in the library (replaces existing entry for the same file). Returns the library."""
    video = os.path.abspath(video)
    srt = srt or (os.path.splitext(video)[0] + ".srt")
    title = title or os.path.splitext(os.path.basename(video))[0]
    data = load_library()
    data["lectures"] = [l for l in data["lectures"] if l.get("video") != video]
    data["lectures"].insert(0, {
        "video": video, "srt": srt, "course": course or "",
        "title": title, "added": time.time(), "viewed": False,
    })
    if course and course not in data["courses"]:
        data["courses"].append(course)
    return save_library(data)


def remove_lecture(video):
    video = os.path.abspath(video)
    data = load_library()
    data["lectures"] = [l for l in data["lectures"] if l.get("video") != video]
    return save_library(data)


def set_lecture_course(video, course):
    video = os.path.abspath(video)
    data = load_library()
    for l in data["lectures"]:
        if l.get("video") == video:
            l["course"] = course or ""
    if course and course not in data["courses"]:
        data["courses"].append(course)
    return save_library(data)


def rename_lecture(video, title):
    """Change display name in the library (does not touch the file on disk — safe during playback)."""
    video = os.path.abspath(video)
    title = (title or "").strip()
    if not title:
        return load_library()
    data = load_library()
    for l in data["lectures"]:
        if l.get("video") == video:
            l["title"] = title
    return save_library(data)


def viewer_path(video):
    """Path to the lecture's standalone HTML player (created at transcription time)."""
    folder = os.path.dirname(video)
    base = os.path.splitext(os.path.basename(video))[0]
    return os.path.join(folder, base + " — כתוביות.html")


def _parse_ts(s):
    s = s.strip().replace(",", ".")
    h, m, rest = s.split(":")
    return int(h) * 3600 + int(m) * 60 + float(rest)


def parse_srt(srt):
    """Read an SRT file back into a list of cues (for opening a saved lecture in the player)."""
    cues = []
    try:
        with open(srt, encoding="utf-8") as f:
            content = f.read()
    except Exception:
        return cues
    for block in re.split(r"\n\s*\n", content.strip()):
        lines = [l for l in block.splitlines() if l.strip()]
        tline = next((l for l in lines if "-->" in l), None)
        if not tline:
            continue
        try:
            a, b = tline.split("-->")
            start, end = _parse_ts(a), _parse_ts(b)
        except Exception:
            continue
        text = " ".join(lines[lines.index(tline) + 1:]).strip()
        if text:
            cues.append({"start": round(start, 3), "end": round(end, 3), "text": text})
    return cues


def open_lecture(video):
    """Returns {video, cues, srt} for a saved lecture — reads SRT from disk. Marks as viewed."""
    video = os.path.abspath(video)
    data = load_library()
    lec = next((l for l in data["lectures"] if l.get("video") == video), None)
    srt = (lec or {}).get("srt") or (os.path.splitext(video)[0] + ".srt")
    if lec and not lec.get("viewed"):
        lec["viewed"] = True
        save_library(data)
    return {"video": video, "cues": parse_srt(srt), "srt": srt}


def search_library(query):
    """Search text across all SRTs in the library. Returns [{video, title, course, hits:[cue,…]}, …]."""
    query = (query or "").strip().lower()
    if not query:
        return []
    data = load_library()
    results = []
    for lec in data["lectures"]:
        cues = parse_srt(lec.get("srt") or "")
        hits = [c for c in cues if query in c["text"].lower()]
        if hits:
            results.append({
                "video": lec["video"],
                "title": lec.get("title") or os.path.splitext(os.path.basename(lec["video"]))[0],
                "course": lec.get("course") or "",
                "hits": hits,
            })
    return results


def download(url, on_progress=None):
    """Download a video from a URL (yt-dlp). Returns the path of the downloaded file.

    Tries with browser cookies first (for Moodle behind login); falls back to
    no cookies for public URLs (YouTube etc.).
    """
    from yt_dlp import YoutubeDL

    outdir = os.path.join(os.path.expanduser("~"), "Videos", "Subtitle Sidekick")
    os.makedirs(outdir, exist_ok=True)
    dl_path = None

    def hook(d):
        nonlocal dl_path
        if d.get("status") == "downloading":
            total = d.get("total_bytes") or d.get("total_bytes_estimate") or 0
            done = d.get("downloaded_bytes") or 0
            pct = int(done / total * 100) if total else 0
            if on_progress:
                on_progress({"percent": pct, "status": "downloading"})
        elif d.get("status") == "finished":
            dl_path = d.get("filename")
            if on_progress:
                on_progress({"percent": 100, "status": "finished"})

    base_opts = {
        "outtmpl": os.path.join(outdir, "%(title).80s.%(ext)s"),
        "format": "best",          # single file (video+audio) — no ffmpeg merge needed
        "progress_hooks": [hook],
        "noplaylist": True,
        "quiet": True,
        "no_warnings": True,
    }

    def run(with_cookies):
        opts = dict(base_opts)
        if with_cookies:
            opts["cookiesfrombrowser"] = ("chrome",)
        with YoutubeDL(opts) as ydl:
            ydl.extract_info(url, download=True)

    # 1) try with cookies (for Moodle behind login)
    cookie_err = None
    try:
        run(True)
        return dl_path
    except Exception as e:  # noqa: BLE001
        cookie_err = str(e)
        dl_path = None

    # 2) fall back to no cookies (for public URLs like YouTube)
    try:
        run(False)
        return dl_path
    except Exception as e2:  # noqa: BLE001
        msg = str(e2)
        low = msg.lower()
        # (a) URL is a player page (e.g. Moodle video page), not a direct video — nothing to extract
        if "unsupported url" in low or "no video" in low or "unable to extract" in low:
            raise RuntimeError(
                "הקישור הזה הוא דף נגן (למשל דף וידאו של Moodle), לא קובץ וידאו ישיר — "
                "אי אפשר להוריד אותו אוטומטית. הורד את הסרטון ידנית וגרור את הקובץ לכאן, "
                "או מצא את כתובת הווידאו הישירה (mp4/m3u8) והדבק אותה.")
        # (b) cookie read error (Chrome open/locked) — needed for Moodle login
        if "could not copy" in cookie_err.lower() or "cookie" in cookie_err.lower():
            raise RuntimeError(
                "ל-Moodle צריך להתחבר בכרום ואז לסגור אותו לגמרי (כדי שאפשר יהיה לקרוא את "
                "ההתחברות), ולנסות שוב.")
        raise RuntimeError(msg)


def _split_words(words):
    """Split a segment's words into short cues of MAX_CUE_WORDS, with word-level timing."""
    out = []
    for i in range(0, len(words), MAX_CUE_WORDS):
        chunk = words[i:i + MAX_CUE_WORDS]
        text = " ".join(w.word.strip() for w in chunk).strip()
        if text:
            out.append({"start": round(chunk[0].start, 3), "end": round(chunk[-1].end, 3), "text": text})
    return out


def transcribe(video_path, fast=False, on_progress=None, cloud=None, pause_check=None):
    """Transcribe a file → SRT + cues. on_progress(dict) is called throughout.

    cloud={"endpoint_url":..., "api_key":...} routes to an external server (cloud_backend)
    instead of the local model. The rest of the app always receives the same result structure.
    pause_check(): optional callable invoked each segment; it may block while paused and
    should return the number of seconds it blocked (for ETA accounting).
    """
    if cloud:
        import cloud_backend
        return cloud_backend.transcribe_remote(
            video_path, cloud.get("endpoint_url", ""), cloud.get("api_key", ""), on_progress)

    global _model, _model_name, _device

    def emit(**kw):
        if on_progress:
            kw.setdefault("device", _device)
            on_progress(kw)

    name = MODEL_FAST if fast else MODEL_ACCURATE
    if _model is None or _model_name != name:
        emit(stage="extract", percent=0, eta=None, elapsed=0, loading=True)
        _model, _device = _load_model(name)
        _model_name = name
    emit(stage="extract", percent=100, eta=None, elapsed=0)

    emit(stage="transcribe", percent=0, eta=None, elapsed=0)
    segments, info = _model.transcribe(
        video_path, language="he", vad_filter=True,
        beam_size=1 if fast else 5,
        word_timestamps=True,
        vad_parameters={"min_silence_duration_ms": 500})
    dur = getattr(info, "duration", 0) or 0

    base = os.path.splitext(video_path)[0]
    srt = base + ".srt"
    cues = []
    t0 = time.time()
    paused_for = 0.0
    for seg in segments:
        if pause_check:                  # cooperative pause — blocks here while paused
            paused_for += pause_check() or 0.0
        words = getattr(seg, "words", None) or []
        sub = _split_words(words) if words else (
            [{"start": round(seg.start, 3), "end": round(seg.end, 3), "text": seg.text.strip()}]
            if seg.text.strip() else []
        )
        cues.extend(sub)
        if sub and dur:
            elapsed = time.time() - t0 - paused_for
            rate = seg.end / elapsed if elapsed > 0 else 0
            eta = (dur - seg.end) / rate if rate > 0 else 0
            emit(stage="transcribe", percent=min(99, int(seg.end / dur * 100)),
                 eta=eta, elapsed=elapsed, line=sub[-1]["text"])

    write_srt(srt, cues)
    emit(stage="sync", percent=99, eta=0, elapsed=time.time() - t0 - paused_for)
    viewer = make_viewer(video_path, cues)
    return {"srt": srt, "cues": cues, "viewer": viewer, "count": len(cues), "video": video_path}
